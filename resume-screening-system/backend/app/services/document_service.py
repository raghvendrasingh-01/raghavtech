"""Résumé document text-extraction dispatch.

Turns an uploaded résumé (raw bytes) into clean plain text regardless of format.
Supported formats:

* **PDF**  – delegated to :mod:`app.services.pdf_service` (pdfminer.six).
* **DOCX** – Microsoft Word ``.docx`` via :mod:`python-docx`.

The correct extractor is chosen by sniffing the file's magic bytes (a PDF starts
with ``%PDF-``; a DOCX is a ZIP container starting with ``PK\\x03\\x04``) rather
than trusting the client-provided content type, which is often wrong or absent.
Both extractors raise :class:`~app.exceptions.PDFExtractionError` on failure so
the route layer can handle a single error type.
"""

from __future__ import annotations

import io
import logging

from app.config import get_settings
from app.exceptions import PDFExtractionError
from app.services.pdf_service import (
    _normalise_whitespace,
    extract_text_from_pdf_bytes,
    looks_like_pdf,
)

logger = logging.getLogger(__name__)

# ZIP local-file-header signature — the container format used by .docx files.
_ZIP_MAGIC = b"PK\x03\x04"


def looks_like_docx(data: bytes) -> bool:
    """Return ``True`` if ``data`` looks like a ZIP/OOXML (.docx) container."""
    return data[:4] == _ZIP_MAGIC


def extract_text_from_docx_bytes(data: bytes) -> str:
    """Extract normalised plain text from in-memory DOCX bytes.

    Args:
        data: The raw bytes of an uploaded ``.docx`` file.

    Returns:
        The extracted, whitespace-normalised text (paragraphs plus table cells).

    Raises:
        PDFExtractionError: If the bytes are empty, are not a valid DOCX, cannot
            be parsed, or contain no extractable text.
    """
    if not data:
        raise PDFExtractionError("Uploaded file is empty.")
    if not looks_like_docx(data):
        raise PDFExtractionError(
            "Uploaded file does not appear to be a valid DOCX."
        )

    try:
        from docx import Document

        document = Document(io.BytesIO(data))
    except Exception as exc:  # python-docx raises a variety of errors
        logger.warning("DOCX parse error: %s", exc)
        raise PDFExtractionError(
            "Could not parse the DOCX. The file may be corrupt or malformed."
        ) from exc

    parts: list[str] = [p.text for p in document.paragraphs]
    # Include table cell text — many résumés lay content out in tables.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    text = _normalise_whitespace("\n".join(parts))
    if not text:
        raise PDFExtractionError("No extractable text found in the DOCX file.")

    max_chars = get_settings().MAX_TEXT_CHARS
    if len(text) > max_chars:
        logger.info(
            "Truncating extracted text from %d to %d chars", len(text), max_chars
        )
        text = text[:max_chars]
    return text


def extract_text_from_upload(data: bytes, filename: str | None = None) -> str:
    """Extract text from an uploaded résumé, dispatching by content sniffing.

    Args:
        data: Raw uploaded file bytes.
        filename: Original filename (used only as a fallback hint).

    Returns:
        Extracted, normalised plain text.

    Raises:
        PDFExtractionError: If the format is unsupported or extraction fails.
    """
    if not data:
        raise PDFExtractionError("Uploaded file is empty.")

    if looks_like_pdf(data):
        return extract_text_from_pdf_bytes(data)
    if looks_like_docx(data):
        return extract_text_from_docx_bytes(data)

    # Last-ditch: fall back on the filename extension for the error message.
    ext = (filename or "").lower().rsplit(".", 1)[-1] if filename else ""
    raise PDFExtractionError(
        f"Unsupported résumé format ('{ext or 'unknown'}'). "
        "Please upload a PDF or DOCX file."
    )
