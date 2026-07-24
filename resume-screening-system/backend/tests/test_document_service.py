"""Tests for the format-dispatching document extraction service.

Covers DOCX extraction end-to-end (built in-memory with python-docx, no binary
fixtures on disk) plus the magic-byte dispatch in ``extract_text_from_upload``.
"""

import io

import pytest

from app.exceptions import PDFExtractionError
from app.services.document_service import (
    extract_text_from_docx_bytes,
    extract_text_from_upload,
    looks_like_docx,
)


def make_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    """Build a valid in-memory .docx containing the given paragraphs/table."""
    from docx import Document

    document = Document()
    for para in paragraphs:
        document.add_paragraph(para)
    if table:
        rows, cols = len(table), len(table[0])
        tbl = document.add_table(rows=rows, cols=cols)
        for r, row in enumerate(table):
            for c, cell_text in enumerate(row):
                tbl.rows[r].cells[c].text = cell_text
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_looks_like_docx_true_for_zip_container():
    data = make_docx(["Hello"])
    assert looks_like_docx(data) is True


def test_looks_like_docx_false_for_pdf(sample_resume_pdf):
    assert looks_like_docx(sample_resume_pdf) is False


def test_extracts_paragraph_text():
    data = make_docx(
        [
            "Jane Doe - Senior Python Engineer",
            "Skills: Docker, PostgreSQL, AWS.",
        ]
    )
    text = extract_text_from_docx_bytes(data)
    assert "Python" in text
    assert "Docker" in text


def test_extracts_table_cell_text():
    data = make_docx(
        ["Experience:"],
        table=[["Kubernetes", "GraphQL"], ["Kafka", "Terraform"]],
    )
    text = extract_text_from_docx_bytes(data)
    assert "Kubernetes" in text
    assert "Terraform" in text


def test_rejects_empty_bytes():
    with pytest.raises(PDFExtractionError):
        extract_text_from_docx_bytes(b"")


def test_rejects_non_docx_bytes():
    with pytest.raises(PDFExtractionError):
        extract_text_from_docx_bytes(b"not a zip at all")


def test_upload_dispatch_pdf(sample_resume_pdf):
    text = extract_text_from_upload(sample_resume_pdf, "resume.pdf")
    assert "FastAPI" in text


def test_upload_dispatch_docx():
    data = make_docx(["Backend engineer with FastAPI and Redis."])
    text = extract_text_from_upload(data, "resume.docx")
    assert "Redis" in text


def test_upload_unsupported_format_raises():
    with pytest.raises(PDFExtractionError):
        extract_text_from_upload(b"\x89PNG\r\n\x1a\n plain image bytes", "img.png")
